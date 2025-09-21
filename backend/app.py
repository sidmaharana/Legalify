import os
import json
import pickle
import base64
import tempfile
from typing import List

from pydantic import BaseModel, Field

from flask import Flask, jsonify, request, send_from_directory
from flask_cors import CORS
from dotenv import load_dotenv

import vertexai
from pypdf import PdfReader
from docx import Document
import pytesseract
from pdf2image import convert_from_path
from PIL import Image


from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_vertexai import VertexAIEmbeddings, VertexAI
from langchain_community.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain, load_summarize_chain, StuffDocumentsChain, LLMChain
from langchain.prompts import PromptTemplate
from langchain.docstore.document import Document as LangchainDocument
from langchain.memory import ConversationBufferMemory
from langchain_core.output_parsers import JsonOutputParser
from process_documents import find_keywords_with_prerequisites

# Import functions from Legal-demistyfier
import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'Legal-demistyfier')))
from analyzer import find_relevant_chunks, get_rag_analysis, discover_concepts_with_ai

# Load environment variables from the .env file
load_dotenv()

# Initialize Flask App
app = Flask(__name__)
CORS(app)

# --- Vertex AI Initialization ---
google_cloud_project = os.getenv("GOOGLE_CLOUD_PROJECT")
google_cloud_location = os.getenv("GOOGLE_CLOUD_LOCATION", None)  # optional
if not google_cloud_project:
    print("Error: GOOGLE_CLOUD_PROJECT environment variable not set. Please set it in your .env file.")
else:
    # If you have location set, include it
    if google_cloud_location:
        vertexai.init(project=google_cloud_project, location=google_cloud_location)
        print(f"Vertex AI initialized for project: {google_cloud_project} (location: {google_cloud_location})")
    else:
        vertexai.init(project=google_cloud_project)
        print(f"Vertex AI initialized for project: {google_cloud_project}")

# Create a Vertex LLM instance (model name may vary depending on availability)
llm = VertexAI(model_name="gemini-2.5-flash")

# --- Load Legal Keywords ---
LEGAL_KEYWORDS = []
try:
    with open("legal_keywords.json", 'r', encoding='utf-8') as f:
        LEGAL_KEYWORDS = json.load(f)
    print(f"Loaded {len(LEGAL_KEYWORDS)} legal keywords from legal_keywords.json")
except FileNotFoundError:
    print("Warning: legal_keywords.json not found. Keyword analysis will not be performed.")
except json.JSONDecodeError:
    print("Error: Could not decode legal_keywords.json. Please check its format.")

# --- Load RAG Vector Index ---
RAG_INDEXED_DATA = []
RAG_INDEX_FILE = os.path.join(os.path.dirname(__file__), '..', 'Legal-demistyfier', 'vector_index.pkl')
try:
    with open(RAG_INDEX_FILE, 'rb') as f:
        RAG_INDEXED_DATA = pickle.load(f)
    print(f"Loaded RAG vector index with {len(RAG_INDEXED_DATA)} entries.")
    print(f"Type of RAG_INDEXED_DATA: {type(RAG_INDEXED_DATA)}")
    print(f"Length of RAG_INDEXED_DATA: {len(RAG_INDEXED_DATA)}")
except FileNotFoundError:
    print(f"Warning: RAG vector index file {RAG_INDEX_FILE} not found. RAG analysis will not be performed.")
except Exception as e:
    print(f"Error loading RAG vector index: {e}")

# --- Pydantic Models for Structured Output ---
class GlossaryTerm(BaseModel):
    """A single term and its definition."""
    term: str = Field(description="The legal jargon or complex phrase identified in the text.")
    definition: str = Field(description="A simple, one-sentence definition of the term in plain English.")

class Glossary(BaseModel):
    """A list of glossary terms."""
    glossary: List[GlossaryTerm] = Field(description="A list of identified legal terms and their definitions.")

# --- Helper Functions ---
def _extract_text_with_ocr(image) -> str:
    """Extracts text from a PIL Image using Tesseract OCR."""
    try:
        return pytesseract.image_to_string(image)
    except pytesseract.TesseractNotFoundError:
        raise RuntimeError(
            "Tesseract is not installed or not in your PATH. Please install it."
        )
    except Exception as e:
        print(f"Error during OCR: {e}")
        return ""

def extract_text_from_file(file) -> List[LangchainDocument]:
    """
    Extracts text content from various document formats (PDF, DOCX, TXT, PNG, JPG, JPEG).
    Performs OCR on PDFs with little text and on image files.
    Returns a list of LangchainDocument objects with metadata.
    """
    file_extension = os.path.splitext(file.filename)[1].lower()
    documents = []

    if file_extension == ".pdf":
        reader = PdfReader(file)
        full_text = []
        for i, page in enumerate(reader.pages):
            page_content = page.extract_text() or ""
            full_text.append(page_content)
            documents.append(LangchainDocument(page_content=page_content, metadata={'source': f'Page {i+1}'}))

        # Always attempt OCR and compare with initial extraction.
        print("Attempting OCR on PDF...")
        file.seek(0)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf.write(file.read())
            temp_pdf_path = temp_pdf.name

        try:
            poppler_bin_path = os.getenv("POPPLER_PATH")
            if poppler_bin_path:
                images = convert_from_path(temp_pdf_path, poppler_path=poppler_bin_path)
            else:
                # Fallback if POPPLER_PATH is not set (e.g., for local development without explicit path)
                print("Warning: POPPLER_PATH environment variable not set. OCR might fail on Render.")
                images = convert_from_path(temp_pdf_path) # This might fail if Poppler is not in system PATH
            ocr_documents = []
            for i, image in enumerate(images):
                ocr_text = _extract_text_with_ocr(image)
                ocr_documents.append(LangchainDocument(page_content=ocr_text, metadata={'source': f'Page {i+1} (OCR)'}))
            
            # Replace original documents with OCR results if OCR found more text
            if len("\n".join([d.page_content for d in ocr_documents]).strip()) > len("\n".join(full_text).strip()):
                documents = ocr_documents
        finally:
            os.remove(temp_pdf_path)

    elif file_extension in [".png", ".jpg", ".jpeg"]:
        image = Image.open(file)
        ocr_text = _extract_text_with_ocr(image)
        documents.append(LangchainDocument(page_content=ocr_text, metadata={'source': 'Image (OCR)'}))

    elif file_extension == ".docx":
        doc = Document(file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        documents.append(LangchainDocument(page_content="\n".join(full_text), metadata={'source': 'Document'}))

    elif file_extension == ".txt":
        text_content = file.read().decode("utf-8")
        documents.append(LangchainDocument(page_content=text_content, metadata={'source': 'Document'}))

    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

    return documents

# --- Flask Routes ---
@app.route('/')
def health_check():
    return jsonify({'status': 'Legalify backend is running'})

@app.route('/api/process', methods=['POST'])
def process_document():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part in the request'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    temp_file_path = None
    try:
        # Save the uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as temp_file:
            file.save(temp_file.name)
            temp_file_path = temp_file.name

        extracted_documents = extract_text_from_file(file)
        full_extracted_text = "\n".join([doc.page_content for doc in extracted_documents])

        # Find keywords in the document using pre-defined legal keywords
        found_keywords_data = find_keywords_with_prerequisites(temp_file_path, LEGAL_KEYWORDS, llm)

        # Summarize the full document (robust to different langchain versions)
        try:
            summary_chain = load_summarize_chain(llm, chain_type="stuff")
            # many versions accept .run or .invoke
            try:
                summary = summary_chain.run(extracted_documents)
            except Exception:
                summary = summary_chain.invoke({"input_documents": extracted_documents})
                # normalize
                if isinstance(summary, dict) and 'output_text' in summary:
                    summary = summary['output_text']
                elif isinstance(summary, dict):
                    # fallback, just stringify
                    summary = str(summary)
        except Exception:
            # Fallback simple prompt-based summarization if load_summarize_chain isn't available
            simple_prompt = PromptTemplate(
                template="Summarize the following document in clear plain English in 4-6 bullet points:\n\n{text}\n",
                input_variables=["text"]
            )
            temp_chain = LLMChain(llm=llm, prompt=simple_prompt)
            summary = temp_chain.run({"text": full_extracted_text})
            if isinstance(summary, dict) and 'output_text' in summary:
                summary = summary['output_text']

        # Split text into chunks for embeddings and FAISS
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        chunks = text_splitter.split_documents(extracted_documents)

        # Create embeddings and vector store
        embeddings_model = VertexAIEmbeddings(model_name="text-embedding-004")
        vector_store = FAISS.from_documents(chunks, embeddings_model)

        # Save FAISS index to temp directory and base64-encode files to return to frontend
        with tempfile.TemporaryDirectory() as temp_dir:
            vector_store.save_local(temp_dir, index_name="my_faiss_index")
            faiss_index_path = os.path.join(temp_dir, "my_faiss_index.faiss")
            pkl_path = os.path.join(temp_dir, "my_faiss_index.pkl")

            # some FAISS versions create other filenames; fallback to any .faiss/.pkl present
            if not os.path.exists(faiss_index_path):
                for fname in os.listdir(temp_dir):
                    if fname.endswith(".faiss"):
                        faiss_index_path = os.path.join(temp_dir, fname)
                        break
            if not os.path.exists(pkl_path):
                for fname in os.listdir(temp_dir):
                    if fname.endswith(".pkl"):
                        pkl_path = os.path.join(temp_dir, fname)
                        break

            with open(faiss_index_path, "rb") as f:
                faiss_index_bytes = f.read()
            with open(pkl_path, "rb") as f:
                pkl_bytes = f.read()

        encoded_faiss_data = {
            "faiss_index": base64.b64encode(faiss_index_bytes).decode('utf-8'),
            "pkl_data": base64.b64encode(pkl_bytes).decode('utf-8')
        }

        # Prepare a robust question generation chain (explicit LLMChain)
        question_gen_template = """
You are a paralegal. Read the following document and generate 3-4 simple, easy-to-understand questions for a non-legal professional about potential risks, obligations, deadlines, or unclear clauses mentioned in the document.
Return ONLY a valid JSON array of strings, where each string is a question.
Document: {document}
"""
        QUESTION_GEN_PROMPT = PromptTemplate(
            template=question_gen_template,
            input_variables=["document"]
        )
        question_generator = LLMChain(llm=llm, prompt=QUESTION_GEN_PROMPT)

        suggested_questions_raw = None
        try:
            suggested_questions_raw = question_generator.run({"document": full_extracted_text})
        except Exception:
            # fallback to invoking if run is not supported
            try:
                suggested_questions_raw = question_generator.invoke({"document": full_extracted_text})
            except Exception as e:
                suggested_questions_raw = "[]"
                print(f"Question generator failed: {e}")

        if isinstance(suggested_questions_raw, str):
            suggested_questions_raw = suggested_questions_raw.strip().replace('```json', '').replace('```', '')
        else:
            # if chain returned dict with text inside
            if isinstance(suggested_questions_raw, dict):
                suggested_questions_raw = suggested_questions_raw.get('output_text', json.dumps([]))
            else:
                suggested_questions_raw = str(suggested_questions_raw)

        suggested_questions = []
        try:
            parsed_questions = json.loads(suggested_questions_raw)
            if isinstance(parsed_questions, list):
                suggested_questions = parsed_questions
        except json.JSONDecodeError:
            # best-effort: try to extract lines that look like questions
            lines = [l.strip() for l in suggested_questions_raw.splitlines() if l.strip().endswith('?')]
            if lines:
                suggested_questions = lines

        return jsonify({
            'extracted_text': full_extracted_text,
            'faiss_index': encoded_faiss_data,
            'summary': summary,
            'suggested_questions': suggested_questions,
            'ai_keyword_analysis': found_keywords_data
        }), 200

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'An unexpected error occurred: {str(e)}'}), 500
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path)

@app.route('/api/rag_analyze', methods=['POST'])
def rag_analyze_document():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part in the request'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    temp_file_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as temp_file:
            file.save(temp_file.name)
            temp_file_path = temp_file.name

        # Extract text from the document
        extracted_documents = extract_text_from_file(file)
        full_extracted_text = "\n".join([doc.page_content for doc in extracted_documents])

        if not full_extracted_text:
            return jsonify({'error': 'Could not extract text from document for RAG analysis.'}), 400

        if not RAG_INDEXED_DATA:
            return jsonify({'error': 'RAG knowledge base not loaded. Please ensure vector_index.pkl exists.'}), 500

        # --- Document Processing for Q&A (copied from /api/process) ---
        # Split text into chunks for embeddings and FAISS
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        chunks = text_splitter.split_documents(extracted_documents)

        # Create embeddings and vector store
        embeddings_model = VertexAIEmbeddings(model_name="text-embedding-004")
        vector_store = FAISS.from_documents(chunks, embeddings_model)

        # Save FAISS index to temp directory and base64-encode files to return to frontend
        with tempfile.TemporaryDirectory() as temp_dir:
            vector_store.save_local(temp_dir, index_name="my_faiss_index")
            faiss_index_path = os.path.join(temp_dir, "my_faiss_index.faiss")
            pkl_path = os.path.join(temp_dir, "my_faiss_index.pkl")

            # some FAISS versions create other filenames; fallback to any .faiss/.pkl present
            if not os.path.exists(faiss_index_path):
                for fname in os.listdir(temp_dir):
                    if fname.endswith(".faiss"):
                        faiss_index_path = os.path.join(temp_dir, fname)
                        break
            if not os.path.exists(pkl_path):
                for fname in os.listdir(temp_dir):
                    if fname.endswith(".pkl"):
                        pkl_path = os.path.join(temp_dir, fname)
                        break

            with open(faiss_index_path, "rb") as f:
                faiss_index_bytes = f.read()
            with open(pkl_path, "rb") as f:
                pkl_bytes = f.read()

        encoded_faiss_data = {
            "faiss_index": base64.b64encode(faiss_index_bytes).decode('utf-8'),
            "pkl_data": base64.b64encode(pkl_bytes).decode('utf-8')
        }
        # --- End Document Processing for Q&A ---

        # Discover concepts using Vertex AI
        search_queries = discover_concepts_with_ai(full_extracted_text, llm)
        if isinstance(search_queries, str):
            return jsonify({'error': f'Error discovering concepts: {search_queries}'}), 500

        rag_analysis_results = []
        for query in search_queries:
            # Retrieve relevant knowledge
            relevant_chunks = find_relevant_chunks(query, RAG_INDEXED_DATA, embeddings_model)
            retrieved_knowledge = "\n\n".join([chunk['text'] for chunk in relevant_chunks])

            # Generate RAG analysis
            analysis = get_rag_analysis(query, retrieved_knowledge, llm)
            rag_analysis_results.append({
                "query": query,
                "analysis": analysis
            })

        # Prepare a robust question generation chain (explicit LLMChain)
        question_gen_template = """
You are a paralegal. Read the following document and generate 3-4 simple, easy-to-understand questions for a non-legal professional about potential risks, obligations, deadlines, or unclear clauses mentioned in the document.
Return ONLY a valid JSON array of strings, where each string is a question.
Document: {document}
"""
        QUESTION_GEN_PROMPT = PromptTemplate(
            template=question_gen_template,
            input_variables=["document"]
        )
        question_generator = LLMChain(llm=llm, prompt=QUESTION_GEN_PROMPT)

        suggested_questions_raw = None
        try:
            suggested_questions_raw = question_generator.run({"document": full_extracted_text})
        except Exception:
            # fallback to invoking if run is not supported
            try:
                suggested_questions_raw = question_generator.invoke({"document": full_extracted_text})
            except Exception as e:
                suggested_questions_raw = "[]"
                print(f"Question generator failed: {e}")

        if isinstance(suggested_questions_raw, str):
            suggested_questions_raw = suggested_questions_raw.strip().replace('```json', '').replace('```', '')
        else:
            # if chain returned dict with text inside
            if isinstance(suggested_questions_raw, dict):
                suggested_questions_raw = suggested_questions_raw.get('output_text', json.dumps([]))
            else:
                suggested_questions_raw = str(suggested_questions_raw)

        suggested_questions = []
        try:
            parsed_questions = json.loads(suggested_questions_raw)
            if isinstance(parsed_questions, list):
                suggested_questions = parsed_questions
        except json.JSONDecodeError:
            # best-effort: try to extract lines that look like questions
            lines = [l.strip() for l in suggested_questions_raw.splitlines() if l.strip().endswith('?')]
            if lines:
                suggested_questions = lines

        return jsonify({
            'extracted_text': full_extracted_text,
            'faiss_index': encoded_faiss_data,
            'rag_analysis': rag_analysis_results,
            'suggested_questions': suggested_questions
        }), 200

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'An unexpected error occurred during RAG analysis: {str(e)}'}), 500
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path)

@app.route('/api/glossary', methods=['POST'])
def get_glossary():
    data = request.get_json()
    text = data.get('extracted_text')
    if not text:
        return jsonify({'error': 'Missing extracted_text'}), 400

    try:
        parser = JsonOutputParser(pydantic_object=Glossary)
        prompt_template = """
Identify all legal jargon, archaic terms, or complex phrases in the provided text chunk.
For each term, provide a simple, one-sentence definition in plain English.
If no jargon is found, return an empty list for the 'glossary'.
{format_instructions}
TEXT CHUNK:
{text}
"""
        prompt = PromptTemplate(
            template=prompt_template,
            input_variables=["text"],
            partial_variables={"format_instructions": parser.get_format_instructions()}
        )

        # Use an explicit chain that returns structured JSON
        chain = LLMChain(llm=llm, prompt=prompt)

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=8000, chunk_overlap=200)
        chunks = text_splitter.split_text(text)

        all_terms = {}
        for chunk in chunks:
            try:
                # call the chain; parse with parser
                raw = None
                try:
                    raw = chain.run({"text": chunk})
                except Exception:
                    raw = chain.invoke({"text": chunk})
                # raw may be JSON text, or dict
                if isinstance(raw, dict):
                    parsed = raw
                else:
                    # use parser to parse string -> pydantic object
                    parsed = parser.parse(raw)
                # parsed should be a dict-like with 'glossary'
                if parsed and isinstance(parsed, dict) and parsed.get('glossary'):
                    for item in parsed['glossary']:
                        # item expected to be dict with 'term' and 'definition'
                        if isinstance(item, dict) and 'term' in item and 'definition' in item:
                            all_terms[item['term']] = item['definition']
                elif hasattr(parsed, 'glossary'):
                    for item in parsed.glossary:
                        all_terms[item.term] = item.definition
            except Exception as e:
                print(f"Could not process a chunk for glossary: {e}")
                continue

        return jsonify(all_terms), 200
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'An error occurred generating glossary: {str(e)}'}), 500

@app.route('/api/ask', methods=['POST'])
def ask_question():
    data = request.get_json()
    question = data.get('question')
    faiss_index_encoded = data.get('faiss_index')
    chat_history_data = data.get('chat_history', [])
    if not question or not faiss_index_encoded:
        return jsonify({'error': 'Missing question or faiss_index'}), 400

    try:
        faiss_index_data = base64.b64decode(faiss_index_encoded["faiss_index"])
        pkl_data = base64.b64decode(faiss_index_encoded["pkl_data"])

        with tempfile.TemporaryDirectory() as temp_dir:
            with open(os.path.join(temp_dir, "my_faiss_index.faiss"), "wb") as f:
                f.write(faiss_index_data)
            with open(os.path.join(temp_dir, "my_faiss_index.pkl"), "wb") as f:
                f.write(pkl_data)

            embeddings_model = VertexAIEmbeddings(model_name="text-embedding-004")
            vector_store = FAISS.load_local(temp_dir, embeddings_model, index_name="my_faiss_index", allow_dangerous_deserialization=True)

        # Recreate memory and seed it with previous chat messages
        memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
        # populate memory safely
        for message in chat_history_data:
            try:
                if message.get('type') == 'human':
                    memory.chat_memory.add_user_message(message.get('content'))
                elif message.get('type') == 'ai':
                    memory.chat_memory.add_ai_message(message.get('content'))
            except Exception:
                # skip malformed history entries
                continue

        # RAG prompt to format answer nicely
        qa_llm_prompt = PromptTemplate(
            template="""
Use the following pieces of context to answer the question at the end.
If you don't know the answer, just say that you don't know, don't try to make up an answer.
Format the entire response using Github-flavored Markdown for clarity. Use headings, bullet points, and bold text where appropriate.

Context: {context}
Question: {question}
Helpful Answer:
""",
            input_variables=["context", "question"]
        )

        llm_chain = LLMChain(llm=llm, prompt=qa_llm_prompt)
        combine_docs_chain = StuffDocumentsChain(
            llm_chain=llm_chain,
            document_variable_name="context"
        )

        # For the question generator used by ConversationalRetrievalChain, define a simple LLMChain
        question_gen_prompt = PromptTemplate(
            template="Rephrase the user's question for retrieval: {question}",
            input_variables=["question"]
        )
        question_generator = LLMChain(llm=llm, prompt=question_gen_prompt)

        qa_chain = ConversationalRetrievalChain(
            retriever=vector_store.as_retriever(),
            memory=memory,
            combine_docs_chain=combine_docs_chain,
            question_generator=question_generator
        )

        # Run the QA chain (robust to .run/.invoke differences)
        try:
            result = qa_chain.run({"question": question})
        except Exception:
            try:
                result = qa_chain.invoke({"question": question})
            except Exception as e:
                # last resort: manually retrieve top docs and call combine_docs_chain
                retriever = vector_store.as_retriever()
                docs = retriever.get_relevant_documents(question)
                # combine_docs_chain expects documents & metadata depending on version
                try:
                    result = combine_docs_chain.run(docs)
                except Exception:
                    result = {"answer": "Could not run QA pipeline: " + str(e)}

        # Normalize result: it might be a string or dict with keys
        answer_text = None
        sources = []
        if isinstance(result, str):
            answer_text = result
        elif isinstance(result, dict):
            # pick common keys
            answer_text = result.get('answer') or result.get('output_text') or result.get('text') or json.dumps(result)
            # Grab source_documents if present
            if 'source_documents' in result and isinstance(result['source_documents'], list):
                for doc in result['source_documents']:
                    if getattr(doc, 'metadata', None) and doc.metadata.get('source'):
                        sources.append({'source': doc.metadata.get('source')})
            elif 'sources' in result and isinstance(result['sources'], list):
                sources = result['sources']
        else:
            answer_text = str(result)

        # Build updated chat history from memory
        updated_chat_history = []
        try:
            for msg in memory.chat_memory.messages:
                # some message objects are simple dicts, others are Message objects
                if isinstance(msg, dict):
                    updated_chat_history.append(msg)
                else:
                    content = getattr(msg, "content", None)
                    role = getattr(msg, "type", None) or getattr(msg, "role", None)
                    if role in ("human", "user"):
                        updated_chat_history.append({'type': 'human', 'content': content})
                    elif role in ("ai", "assistant"):
                        updated_chat_history.append({'type': 'ai', 'content': content})
                    else:
                        # default to 'ai' if unclear
                        updated_chat_history.append({'type': 'ai', 'content': content})
        except Exception:
            pass

        return jsonify({'answer': answer_text, 'chat_history': updated_chat_history, 'sources': sources}), 200

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'An error occurred during Q&A: {str(e)}'}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    # Set debug=False for production; enable as needed
    app.run(debug=False, host='0.0.0.0', port=port)

# --- Frontend Serving Routes ---
@app.route('/')
def serve_frontend():
    return send_from_directory('../frontend', 'index.html')

@app.route('/<path:path>')
def serve_frontend_files(path):
    return send_from_directory('../frontend', path)
